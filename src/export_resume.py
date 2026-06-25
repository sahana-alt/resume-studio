#!/usr/bin/env python3
"""Export a tailored resume Markdown file to DOCX and PDF.

The layout follows the provided Jake Gutierrez-style LaTeX resume structure:
compact one-page target, centered heading, section rules, bold heading/date
rows, compact bullets, and technical skills with bold subheadings.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def applicant_resume_lines(markdown: str) -> list[str]:
    lines = []
    for line in markdown.splitlines():
        if line.strip() == "## Truthfulness Notes":
            break
        lines.append(line.rstrip())
    return lines


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in text)


def latex_bold_left_date(text: str) -> str:
    if "\t" not in text:
        return rf"\textbf{{{latex_escape(text)}}}"
    left, right = text.split("\t", 1)
    return rf"\textbf{{{latex_escape(left)}}} & \textbf{{\small {latex_escape(right)}}}"


def build_latex(lines: list[str], out_path: Path) -> None:
    body: list[str] = []
    current_section = ""
    item_list_open = False
    subheading_list_open = False
    contact = ""

    def close_item_list(experience_gap: bool = False, project_gap: bool = False) -> None:
        nonlocal item_list_open
        if item_list_open:
            body.append(r"\resumeItemListEnd")
            if experience_gap:
                body.append("")
                body.append(r"\vspace{1pt}")
                body.append("")
            if project_gap:
                body.append("")
                body.append(r"\vspace{-1pt}")
                body.append("")
            item_list_open = False

    def close_subheading_list() -> None:
        nonlocal subheading_list_open
        close_item_list()
        if subheading_list_open:
            body.append(r"\resumeSubHeadingListEnd")
            subheading_list_open = False

    def ensure_subheading_list() -> None:
        nonlocal subheading_list_open
        if not subheading_list_open:
            body.append(r"\resumeSubHeadingListStart")
            subheading_list_open = True

    previous_was_bullet = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            body.append(r"\begin{center}")
            body.append(rf"    {{\Huge \scshape {latex_escape(line[2:])}}} \\ \vspace{{1pt}}")
            continue
        if "@" in line and "|" in line:
            contact = line
            parts = [part.strip() for part in line.split("|")]
            contact_parts = [latex_escape(part) for part in parts]
            body.append(rf"    \small {' ~ '.join(contact_parts)}")
            body.append(r"    \vspace{-8pt}")
            body.append(r"\end{center}")
            continue
        if line.startswith("## "):
            close_subheading_list()
            current_section = line[3:].strip().upper()
            body.append(rf"\section{{{latex_escape(current_section.title())}}}")
            if current_section in {"EDUCATION", "WORK EXPERIENCE", "PROJECTS"}:
                ensure_subheading_list()
            continue
        if line.startswith("### "):
            close_item_list(
                experience_gap=current_section == "WORK EXPERIENCE",
                project_gap=current_section == "PROJECTS",
            )
            ensure_subheading_list()
            heading = line[4:]
            if current_section == "PROJECTS":
                title, date = (heading.split("\t", 1) + [""])[:2] if "\t" in heading else (heading, "")
                body.append(rf"\resumeProjectHeading{{\textbf{{{latex_escape(title)}}}}}{{{latex_escape(date)}}}")
            else:
                title, date = (heading.split("\t", 1) + [""])[:2] if "\t" in heading else (heading, "")
                if " - " in title:
                    company, role = title.rsplit(" - ", 1)
                else:
                    company, role = title, ""
                body.append(rf"\resumeSubheading{{{latex_escape(role or company)}}}{{}}{{{latex_escape(company if role else '')}}}{{{latex_escape(date)}}}")
            body.append(r"\resumeItemListStart")
            item_list_open = True
            continue
        if line.startswith("- "):
            if current_section in {"PUBLICATIONS", "AWARDS AND RECOGNITION"}:
                close_item_list()
                body.append(rf"\small{{{latex_escape(line[2:])}}}\\[-2pt]")
            else:
                if not item_list_open:
                    body.append(r"\resumeItemListStart")
                    item_list_open = True
                body.append(rf"\resumeItem{{{latex_escape(line[2:])}}}")
            continue
        if "\t" in line and current_section == "EDUCATION":
            ensure_subheading_list()
            left, right = line.split("\t", 1)
            pending_degree = ""
            # The next plain line is handled by the main loop, but this marker
            # keeps education rows spacious enough to match the LaTeX template.
            body.append(r"\item")
            body.append(r"\begin{tabular*}{1.0\textwidth}[t]{l@{\extracolsep{\fill}}r}")
            body.append(rf"  \textbf{{{latex_escape(left)}}} & \textbf{{\small {latex_escape(right)}}} \\")
            body.append(r"\end{tabular*}\vspace{-3pt}")
            continue
        if current_section == "TECHNICAL SKILLS":
            if ":" in line:
                left, right = line.split(":", 1)
                body.append(rf"\textbf{{{latex_escape(left)}}}{{: {latex_escape(right.strip())}}} \\")
            else:
                body.append(latex_escape(line) + r" \\")
            continue
        close_item_list()
        if current_section == "EDUCATION":
            body.append(rf"\textit{{\small {latex_escape(line)}}}\\[2pt]")
        elif current_section == "SUMMARY":
            body.append(rf"\small{{{latex_escape(line)}}}\\[-2pt]")
        else:
            body.append(latex_escape(line) + r"\\")

    close_subheading_list()

    latex = "\n".join([
        r"%-------------------------",
        r"% Resume in Latex",
        r"% Author : Jake Gutierrez",
        r"% Based off of: https://github.com/sb2nov/resume",
        r"% License : MIT",
        r"%------------------------",
        r"\documentclass[letterpaper,11pt]{article}",
        r"\usepackage{latexsym}",
        r"\usepackage[empty]{fullpage}",
        r"\usepackage{titlesec}",
        r"\usepackage{marvosym}",
        r"\usepackage[usenames,dvipsnames]{color}",
        r"\usepackage{verbatim}",
        r"\usepackage{enumitem}",
        r"\usepackage[hidelinks]{hyperref}",
        r"\usepackage{fancyhdr}",
        r"\usepackage[english]{babel}",
        r"\usepackage{tabularx}",
        r"\usepackage{fontawesome5}",
        r"\usepackage{multicol}",
        r"\setlength{\multicolsep}{-3.0pt}",
        r"\setlength{\columnsep}{-1pt}",
        r"\IfFileExists{glyphtounicode.tex}{\input{glyphtounicode}}{}",
        r"% \usepackage[inline]{enumitem}",
        r"\makeatletter",
        r"\newcommand{\inlineitem}[1][]{%",
        r"\ifnum\enit@type=\tw@",
        r"    {\descriptionlabel{#1}}",
        r"  \hspace{\labelsep}%",
        r"\else",
        r"  \ifnum\enit@type=\z@",
        r"       \refstepcounter{\@listctr}\fi",
        r"    \quad\@itemlabel\hspace{\labelsep}%",
        r"\fi}",
        r"\makeatother",
        r"\usepackage[svgnames]{xcolor}",
        r"\usepackage{hyperref}",
        r"\hypersetup{",
        r"    colorlinks=true,",
        r"    linkcolor=blue,",
        r"    filecolor=magenta,",
        r"    urlcolor=MidnightBlue,",
        r"    pdftitle={Resume},",
        r"    pdfpagemode=FullScreen,",
        r"    }",
        r"\urlstyle{same}",
        r"\pagestyle{fancy}",
        r"\fancyhf{}",
        r"\fancyfoot{}",
        r"\renewcommand{\headrulewidth}{0pt}",
        r"\renewcommand{\footrulewidth}{0pt}",
        r"% Adjust margins",
        r"\addtolength{\oddsidemargin}{-0.6in}",
        r"\addtolength{\evensidemargin}{-0.5in}",
        r"\addtolength{\textwidth}{1.19in}",
        r"\addtolength{\topmargin}{-.7in}",
        r"\addtolength{\textheight}{1.4in}",
        r"\urlstyle{same}",
        r"\raggedbottom",
        r"\raggedright",
        r"\setlength{\tabcolsep}{0in}",
        r"% Sections formatting",
        r"\titleformat{\section}{",
        r"  \vspace{-4pt}\scshape\raggedright\large\bfseries",
        r"}{}{0em}{}[\color{black}\titlerule \vspace{-5pt}]",
        r"% Ensure that generated pdf is machine readable/ATS parsable",
        r"\pdfgentounicode=1",
        r"%-------------------------",
        r"% Custom commands",
        r"\newcommand{\resumeItem}[1]{",
        r"  \item\small{",
        r"    {#1 \vspace{-2pt}}",
        r"  }",
        r"}",
        r"\newcommand{\classesList}[4]{",
        r"    \item\small{",
        r"        {#1 #2 #3 #4 \vspace{-2pt}}",
        r"  }",
        r"}",
        r"\newcommand{\resumeEduSubheading}[6]{",
        r"  \vspace{-2pt}\item",
        r"    \begin{tabular*}{1.0\textwidth}[t]{l@{\extracolsep{\fill}}r}",
        r"      \textbf{#1} & \textbf{\small #2} \\",
        r"      \textit{\small#3} & \text{\small #4} \\",
        r"      \text{\small#5 :} \text{\small #6} \\",
        r"    \end{tabular*}\vspace{-7pt}",
        r"}",
        r"\newcommand{\resumeSubheading}[4]{",
        r"  \vspace{-2pt}\item",
        r"    \begin{tabular*}{1.0\textwidth}[t]{l@{\extracolsep{\fill}}r}",
        r"      \textbf{#1} & \textbf{\small #2} \\",
        r"      \textit{\small#3} & \textit{\small #4} \\",
        r"    \end{tabular*}\vspace{-7pt}",
        r"}",
        r"\newcommand{\resumeSubSubheading}[2]{",
        r"    \item",
        r"    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}",
        r"      \textit{\small#1} & \textit{\small #2} \\",
        r"    \end{tabular*}\vspace{-7pt}",
        r"}",
        r"\newcommand{\resumeProjectHeading}[2]{",
        r"    \item",
        r"    \begin{tabular*}{1.00\textwidth}{l@{\extracolsep{\fill}}r}",
        r"      \small#1 & \textbf{\small #2}\\",
        r"    \end{tabular*}\vspace{-7pt}",
        r"}",
        r"\newcommand{\resumeProjectHeadingURL}[3]{",
        r"    \item",
        r"    \begin{tabular*}{1.00\textwidth}{l@{\extracolsep{\fill}}r}",
        r"      \small\href{#1}{#2} & \textbf{\small #3}\\",
        r"    \end{tabular*}\vspace{-7pt}",
        r"}",
        r"\newcommand{\resumeProjectHeadingTwo}[2]{",
        r"\textbf{\small#1}",
        r"\hfill",
        r"\textbf{\small#2}",
        r"}",
        r"\newcommand{\resumeSubItem}[1]{\resumeItem{#1}\vspace{-4pt}}",
        r"\renewcommand\labelitemi{$\vcenter{\hbox{\tiny$\bullet$}}$}",
        r"\renewcommand\labelitemii{$\vcenter{\hbox{\tiny$\bullet$}}$}",
        r"\newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.0in, label={}]}\newcommand{\resumeSubHeadingListEnd}{\end{itemize}}",
        r"\newcommand{\resumeItemListStart}{\begin{itemize}}",
        r"\newcommand{\resumeItemListEnd}{\end{itemize}\vspace{-5pt}}",
        r"%-------------------------------------------",
        r"%%%%%%  RESUME STARTS HERE  %%%%%%%%%%%%%%%%%%%%%%%%%%%%",
        r"\begin{document}",
        *body,
        r"\end{document}",
        "",
    ])
    out_path.write_text(latex, encoding="utf-8")


def latex_compiler() -> str | None:
    for name in ("pdflatex", "xelatex", "tectonic"):
        found = shutil.which(name)
        if found:
            return found
    return None


def compile_latex(tex_path: Path, pdf_path: Path) -> bool:
    compiler = latex_compiler()
    if not compiler:
        return False
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        work_tex = tmp_path / tex_path.name
        work_tex.write_text(tex_path.read_text(encoding="utf-8"), encoding="utf-8")
        if Path(compiler).name == "tectonic":
            cmd = [compiler, str(work_tex)]
        else:
            cmd = [compiler, "-interaction=nonstopmode", "-halt-on-error", str(work_tex)]
        result = subprocess.run(cmd, cwd=tmp_path, capture_output=True, text=True, timeout=180)
        if result.returncode != 0:
            log_path = tex_path.with_suffix(".latex.log")
            log_path.write_text(result.stdout + "\n" + result.stderr, encoding="utf-8")
            return False
        built_pdf = tmp_path / tex_path.with_suffix(".pdf").name
        if not built_pdf.exists():
            return False
        pdf_path.write_bytes(built_pdf.read_bytes())
        return True


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = paragraph._p._new_hyperlink()
    hyperlink.set(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
        rel_id,
    )
    run = paragraph.add_run(text)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    run.font.color.rgb = RGBColor(5, 99, 193)
    run.font.underline = True


def build_docx(lines: list[str], out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.30)
    section.bottom_margin = Inches(0.30)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.6)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(10)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(9.5)

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(15)
            previous_was_bullet = False
        elif line.startswith("## "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(line[3:].upper())
            run.bold = True
            run.font.size = Pt(9.6)
            add_bottom_border(paragraph)
            previous_was_bullet = False
        elif line.startswith("### "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
            paragraph.paragraph_format.space_before = Pt(4 if previous_was_bullet else 1)
            paragraph.paragraph_format.space_after = Pt(0)
            text = line[4:]
            if "\t" in text:
                left, right = text.split("\t", 1)
                run = paragraph.add_run(left)
                run.bold = True
                run.font.size = Pt(8.8)
                paragraph.add_run("\t")
                date_run = paragraph.add_run(right)
                date_run.bold = True
                date_run.font.size = Pt(8.6)
            else:
                run = paragraph.add_run(text)
                run.bold = True
                run.font.size = Pt(8.8)
            previous_was_bullet = False
        elif "\t" in line:
            left, right = line.split("\t", 1)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(left)
            run.bold = True
            run.font.size = Pt(8.7)
            paragraph.add_run("\t")
            date_run = paragraph.add_run(right)
            date_run.bold = True
            date_run.font.size = Pt(8.6)
            previous_was_bullet = False
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.16)
            paragraph.paragraph_format.first_line_indent = Inches(-0.10)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.add_run(line[2:]).font.size = Pt(8.45)
            previous_was_bullet = True
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if "@" in line and "|" in line else WD_ALIGN_PARAGRAPH.LEFT
            if ":" in line and not line.startswith("http"):
                left, right = line.split(":", 1)
                run = paragraph.add_run(f"{left}:")
                run.bold = True
                run.font.size = Pt(8.55)
                paragraph.add_run(right).font.size = Pt(8.55)
            else:
                paragraph.add_run(line).font.size = Pt(8.55)
            previous_was_bullet = False

    doc.save(out_path)


def build_cover_letter_docx(markdown: str, profile: dict[str, str], out_path: Path) -> None:
    """Build a restrained one-page business letter from generated cover text."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    name = str(profile.get("name", "Your Name"))
    contact = " | ".join(
        value for value in [
            str(profile.get("email", "")).strip(),
            str(profile.get("phone", "")).strip(),
        ] if value
    )

    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(2)
    name_run = header.add_run(name)
    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(15)

    if contact:
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.paragraph_format.space_after = Pt(18)
        contact_run = contact_paragraph.add_run(contact)
        contact_run.font.name = "Arial"
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(75, 75, 75)

    content = [
        line.strip()
        for line in markdown.splitlines()
        if line.strip() and not line.startswith("# ")
    ]
    for index, line in enumerate(content):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        if index == 0:
            paragraph.paragraph_format.space_after = Pt(14)
        run = paragraph.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(11)

    doc.core_properties.title = "Cover Letter"
    doc.core_properties.subject = "Job application cover letter"
    doc.core_properties.author = name
    doc.save(out_path)


def add_bottom_border(paragraph) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def pdf_heading_with_date(text: str) -> str:
    if "\t" not in text:
        return f"<b>{pdf_escape(text)}</b>"
    left, right = text.split("\t", 1)
    return f"<b>{pdf_escape(left)}</b> {pdf_escape(right)}"


def pdf_heading_table(text: str, style: ParagraphStyle, bold_left: bool = True) -> Table | Paragraph:
    if "\t" not in text:
        return Paragraph(f"<b>{pdf_escape(text)}</b>" if bold_left else pdf_escape(text), style)
    left, right = text.split("\t", 1)
    left_text = f"<b>{pdf_escape(left)}</b>" if bold_left else pdf_escape(left)
    table = Table(
        [[Paragraph(left_text, style), Paragraph(pdf_escape(right), style)]],
        colWidths=[5.55 * inch, 1.35 * inch],
        hAlign="LEFT",
    )
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    return table


def pdf_skill_row(text: str) -> str:
    if ":" not in text:
        return pdf_escape(text)
    left, right = text.split(":", 1)
    return f"<b>{pdf_escape(left)}:</b>{pdf_escape(right)}"


def build_pdf(lines: list[str], out_path: Path) -> None:
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "ResumeNormal",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.25,
        leading=9.15,
        spaceAfter=0.3,
    )
    title = ParagraphStyle(
        "ResumeTitle",
        parent=normal,
        alignment=1,
        fontSize=15,
        leading=16,
        spaceAfter=2,
        textColor=colors.black,
    )
    section = ParagraphStyle(
        "ResumeSection",
        parent=normal,
        fontSize=9.5,
        leading=10.2,
        spaceBefore=3,
        spaceAfter=0.2,
        textColor=colors.black,
        borderWidth=0.5,
        borderColor=colors.black,
        borderPadding=1,
    )
    role = ParagraphStyle(
        "ResumeRole",
        parent=normal,
        fontSize=8.55,
        leading=9.35,
        spaceBefore=1,
        spaceAfter=0,
    )
    bullet = ParagraphStyle(
        "ResumeBullet",
        parent=normal,
        leftIndent=10,
        firstLineIndent=-6,
        bulletIndent=0,
    )
    centered = ParagraphStyle("Centered", parent=normal, alignment=1)

    story = []
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            story.append(Paragraph(f"<b>{pdf_escape(line[2:])}</b>", title))
        elif line.startswith("## "):
            story.append(Paragraph(f"<b>{pdf_escape(line[3:].upper())}</b>", section))
        elif line.startswith("### "):
            story.append(pdf_heading_table(line[4:], role))
        elif "\t" in line:
            story.append(pdf_heading_table(line, normal))
        elif line.startswith("- "):
            story.append(Paragraph(pdf_escape(line[2:]), bullet, bulletText="-"))
        else:
            style = centered if "@" in line and "|" in line else normal
            story.append(Paragraph(pdf_skill_row(line), style))
    story.append(Spacer(1, 0.01 * inch))

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        rightMargin=0.45 * inch,
        leftMargin=0.45 * inch,
        topMargin=0.30 * inch,
        bottomMargin=0.30 * inch,
    )
    doc.build(story)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("markdown", type=Path)
    parser.add_argument("--docx", type=Path)
    parser.add_argument("--pdf", type=Path)
    args = parser.parse_args()

    lines = applicant_resume_lines(args.markdown.read_text(encoding="utf-8"))
    docx_path = args.docx or args.markdown.with_suffix(".docx")
    pdf_path = args.pdf or args.markdown.with_suffix(".pdf")
    tex_path = args.markdown.with_suffix(".tex")
    if pdf_path.exists():
        pdf_path.unlink()
    build_docx(lines, docx_path)
    build_latex(lines, tex_path)
    compiled = compile_latex(tex_path, pdf_path)
    print(f"Wrote {docx_path}")
    print(f"Wrote {tex_path}")
    if compiled:
        print(f"Wrote {pdf_path}")
    else:
        print("Skipped PDF: no LaTeX compiler found or LaTeX compilation failed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
